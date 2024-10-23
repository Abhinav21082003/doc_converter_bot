import streamlit as st
import pdfplumber
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import google.generativeai as genai
from docx import Document
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Hard-coded API key (replace 'your_api_key_here' with your actual API key)
api_key = "AIzaSyBFn62709Z2xFN4TMtRpv9a_HLVFDxBpgg"

# Set the path to the Tesseract executable
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Adjust the path as necessary

def sanitize_text(text):
    return ''.join(c for c in text if c.isprintable())

def wrap_text(text, max_width, font_size):
    """Wrap text to fit within the specified width."""
    wrapped_lines = []
    lines = text.split('\n')
    for line in lines:
        while len(line) > max_width:
            # Find the last space within the maximum width
            split_index = line.rfind(' ', 0, max_width)
            if split_index == -1:  # No space found, break the line at max width
                split_index = max_width
            wrapped_lines.append(line[:split_index])
            line = line[split_index:].strip()
        wrapped_lines.append(line)
    return wrapped_lines

if not api_key:
    st.error("API key not found.")
else:
    # Configure the Gemini API with the API key
    genai.configure(api_key=api_key)

    # Initialize the model
    model = genai.GenerativeModel("gemini-pro")
    chat = model.start_chat(history=[])

    # Streamlit app layout
    st.title("Document Conversion and Chatbot")

    # Upload file
    uploaded_file = st.file_uploader("Upload an image, PDF, text, or Word file", type=["jpg", "jpeg", "png", "pdf", "txt", "docx"])

    if uploaded_file:
        # Ask user to input the desired filename
        filename = st.text_input("Enter the filename (without extension):", "output")

        if uploaded_file.type == "application/pdf":
            # Extract text from PDF
            with pdfplumber.open(uploaded_file) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""

            # Display and download the extracted text
            st.subheader("Extracted Text from PDF")
            st.write(text)

            # Convert extracted text to a text file
            text_file_name = f"{filename}.txt"
            with open(text_file_name, "w", encoding="utf-8") as text_file:
                text_file.write(text)
            st.download_button(label=f"Download Extracted Text as {text_file_name}", data=text, file_name=text_file_name)

            # Convert PDF to Word
            doc = Document()
            sanitized_text = sanitize_text(text)
            doc.add_paragraph(sanitized_text)
            word_file_name = f"{filename}.docx"
            doc.save(word_file_name)
            with open(word_file_name, "rb") as doc_file:
                st.download_button(label=f"Download as Word Document ({word_file_name})", data=doc_file, file_name=word_file_name)

        elif uploaded_file.type == "text/plain":
            # Convert text file to PDF and JPEG
            text = uploaded_file.read().decode("utf-8")
            st.subheader("Uploaded Text")
            st.write(text)

            # Handle unsupported characters by replacing or removing them
            safe_text = sanitize_text(text)

            # Convert text to PDF using reportlab
            pdf_file_name = f"{filename}.pdf"
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            text_object = c.beginText(40, height - 40)
            text_object.setFont("Helvetica", 12)

            wrapped_lines = wrap_text(safe_text, 100, 12)  # Font size is 12
            for line in wrapped_lines:
                text_object.textLine(line)

            c.drawText(text_object)
            c.showPage()
            c.save()
            buffer.seek(0)
            st.download_button(label=f"Download Text as PDF ({pdf_file_name})", data=buffer, file_name=pdf_file_name)

            # Convert text to JPEG
            font = ImageFont.load_default()
            image = Image.new('RGB', (800, 400), color=(255, 255, 255))
            draw = ImageDraw.Draw(image)
            draw.text((10, 10), safe_text, font=font, fill=(0, 0, 0))
            img_buffer = io.BytesIO()
            image.save(img_buffer, format="JPEG")
            st.download_button(label=f"Download Text as JPEG ({filename}.jpeg)", data=img_buffer.getvalue(), file_name=f"{filename}.jpeg")

            # Convert text to Word
            doc = Document()
            doc.add_paragraph(safe_text)
            word_file_name = f"{filename}.docx"
            doc.save(word_file_name)
            with open(word_file_name, "rb") as doc_file:
                st.download_button(label=f"Download Text as Word Document ({word_file_name})", data=doc_file, file_name=word_file_name)

        elif uploaded_file.type in ["image/jpeg", "image/png"]:
            # Extract text from image
            image = Image.open(uploaded_file)
            text = pytesseract.image_to_string(image)
            st.subheader("Extracted Text from Image")
            st.write(text)

            # Convert extracted text to a text file
            text_file_name = f"{filename}.txt"
            sanitized_text = sanitize_text(text)
            with open(text_file_name, "w", encoding="utf-8") as text_file:
                text_file.write(sanitized_text)
            st.download_button(label=f"Download Extracted Text as {text_file_name}", data=sanitized_text, file_name=text_file_name)

            # Convert image text to PDF using reportlab
            pdf_file_name = f"{filename}.pdf"
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            text_object = c.beginText(40, height - 40)

            wrapped_lines = wrap_text(sanitized_text, 100, 12)  # Font size is 12
            for line in wrapped_lines:
                text_object.textLine(line)

            c.drawText(text_object)
            c.showPage()
            c.save()
            buffer.seek(0)
            st.download_button(label=f"Download Image Text as PDF ({pdf_file_name})", data=buffer, file_name=pdf_file_name)

        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Extract text from Word file
            doc = Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            st.subheader("Extracted Text from Word Document")
            st.write(text)

            # Convert Word text to PDF using reportlab
            pdf_file_name = f"{filename}.pdf"
            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)
            width, height = letter
            text_object = c.beginText(40, height - 40)

            wrapped_lines = wrap_text(text, 100, 12)  # Font size is 12
            for line in wrapped_lines:
                text_object.textLine(line)

            c.drawText(text_object)
            c.showPage()
            c.save()
            buffer.seek(0)
            st.download_button(label=f"Download Word Text as PDF ({pdf_file_name})", data=buffer, file_name=pdf_file_name)

        # Chat interaction
        st.subheader("Ask a question based on the extracted text")
        user_question = st.text_input("Your Question:")

        if user_question and (text or sanitized_text):
            try:
                # Send the extracted text and user question to the chatbot
                prompt = f"The following text was extracted from the document:\n\n{text or sanitized_text}\n\nBased on this, the user asks: {user_question}\n\nPlease provide an answer."
                response = chat.send_message(prompt)
                st.write(f"**Bot:** {response.text}")

                # Download the chatbot's response as a PDF
                response_file_name = f"{filename}_response.pdf"
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                width, height = letter
                text_object = c.beginText(40, height - 40)
                text_object.setFont("Helvetica", 12)

                wrapped_response = wrap_text(response.text, 100, 12)  # Font size is 12
                text_object.textLines(["Chatbot Response:"] + wrapped_response)

                c.drawText(text_object)
                c.showPage()
                c.save()
                buffer.seek(0)
                st.download_button(label=f"Download Chatbot Response as PDF ({response_file_name})", data=buffer, file_name=response_file_name)

            except Exception as e:
                st.write(f"Error occurred: {e}")
